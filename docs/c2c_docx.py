#!/usr/bin/env python3
"""Mise en page Camptocamp des guides utilisateur et guides de décision (DOCX + PDF).

Bibliothèque extraite des guides livrés à Stucki Leadership (août 2026) et à
RubixComm (juillet-août 2026). Un générateur de guide l'importe, déclare sa
`Brand` (client, logo, langue), puis enchaîne les briques dans l'ordre du
document : couverture, encadré de synthèse, sections numérotées, captures
légendées, matrices, « bon à savoir », point ouvert, encadré de vérification.

    from c2c_docx import Brand, Guide, Cm, Inches

    brand = Brand(client="Stucki Leadership", client_logo="stucki_leadership_logo.png",
                  kind="User guide", lang="en")
    g = Guide(brand, shots_dir=ROOT / "docs/screenshots", assets_dir=ROOT / "docs/assets")
    g.cover("The two project buttons", "Add Coordination Tasks & Add Insights Task",
            "Stucki Leadership — 28 August 2026")
    g.callout("Both buttons now work for everyone who can open a project…")
    g.heading("1. What the two buttons do")
    g.body("The two buttons sit in the header of every seminar project…")
    g.screenshot("buttons_project_header.png", "The two buttons in the header.")
    g.matrix(["Button", "What it puts in the project"], [("Add Coordination Tasks", "…")])
    g.save(ROOT / "docs/Stucki_Odoo19_Project_Task_Buttons_Guide.docx", pdf=True)

Dépendances : python-docx (DOCX), LibreOffice (`soffice`, conversion PDF),
poppler-utils (`pdftoppm`, relecture visuelle page par page).
"""

from __future__ import annotations

import shutil
import subprocess
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

__all__ = ["Brand", "Guide", "Cm", "Inches", "Pt", "to_pdf", "pdf_pages_to_png"]

HERE = Path(__file__).resolve().parent
DEFAULT_ASSETS = HERE / "assets"

# --- Charte ------------------------------------------------------------------
# L'orange Camptocamp porte les titres, les en-têtes de tableau, les numéros
# d'étape et le fond des encadrés (dans sa version pâle). Tout le reste est gris.
ORANGE = "FF6600"
DARK = "252525"
GREY = "7F8385"
MID_GREY = "D9DCDD"
LIGHT_GREY = "F1F2F2"
LIGHT_ORANGE = "FFF0E6"
GREEN = "138A72"
RED = "C43D35"
WHITE = "FFFFFF"
FONT = "Lato"

# Largeur utile d'une page A4 avec les marges de `setup_page` : 17,7 cm.
CONTENT_WIDTH = Cm(17.7)
SHOT_WIDTH = Inches(6.45)      # capture pleine largeur
SHOT_WIDTH_NARROW = Inches(3.4)  # notification, popup, assistant étroit
SHOT_WIDTH_PAIR = Inches(3.05)   # deux captures côte à côte

LABELS = {
    "fr": {"page": "Page", "copyright": "© Camptocamp | {client} — {subject}"},
    "en": {"page": "Page", "copyright": "© Camptocamp | {client} — {subject}"},
    "de": {"page": "Seite", "copyright": "© Camptocamp | {client} — {subject}"},
}


@dataclass
class Brand:
    """Identité du document : le client, son logo, le type de guide, la langue."""

    client: str
    client_logo: str | None = None        # fichier dans assets_dir (≈ 1,25 cm de large en en-tête)
    kind: str = "Guide utilisateur"       # libellé de l'en-tête : « User guide », « Decision guide »…
    lang: str = "fr"
    subject: str = ""                     # complément du pied de page ; défaut : kind
    accent: str = ORANGE
    font: str = FONT
    author: str = "Camptocamp"
    client_logo_width: object = field(default_factory=lambda: Cm(1.25))
    client_logo_cover_width: object = field(default_factory=lambda: Cm(3.1))


# --- Bas niveau : XML python-docx ------------------------------------------
def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MID_GREY, size=4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tbl_pr.append(borders)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def add_field(paragraph, instruction: str) -> None:
    """Champ Word (PAGE, NUMPAGES…), calculé à l'ouverture et à la conversion PDF."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, text, separate, end))


def style_run(run, size=10, bold=False, color=DARK, italic=False, font=FONT) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_text(cell, text, *, bold=False, color=DARK, size=8.5, align=None, font=FONT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    style_run(p.add_run(str(text)), size=size, bold=bold, color=color, font=font)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_column_widths(table, widths) -> None:
    for column, width in zip(table.columns, widths):
        column.width = width
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width


# --- Le document -------------------------------------------------------------
class Guide:
    """Un guide Camptocamp : A4 portrait, Lato, en-tête et pied de page de marque."""

    def __init__(self, brand: Brand, *, shots_dir: Path, assets_dir: Path | None = None):
        self.brand = brand
        self.shots = Path(shots_dir)
        self.assets = Path(assets_dir) if assets_dir else DEFAULT_ASSETS
        self.doc = Document()
        self._setup_styles()
        self._setup_page(self.doc.sections[0])
        self._setup_header_footer(self.doc.sections[0])
        self.doc.core_properties.author = brand.author

    # -- réglages -------------------------------------------------------------
    def _setup_styles(self) -> None:
        doc, b = self.doc, self.brand
        normal = doc.styles["Normal"]
        normal.font.name = b.font
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), b.font)
        normal.font.size = Pt(9.5)
        normal.font.color.rgb = RGBColor.from_string(DARK)
        normal.paragraph_format.space_after = Pt(5)
        for style_name, size, color in (
            ("Title", 28, b.accent),
            ("Heading 1", 20, b.accent),
            ("Heading 2", 13, DARK),
            ("Heading 3", 10.5, b.accent),
        ):
            style = doc.styles[style_name]
            style.font.name = b.font
            style._element.rPr.rFonts.set(qn("w:eastAsia"), b.font)
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(0 if style_name == "Heading 1" else 5)
            style.paragraph_format.space_after = Pt(6 if style_name == "Heading 1" else 4)
        for style_name in ("List Bullet", "List Bullet 2"):
            style = doc.styles[style_name]
            style.font.name = b.font
            style._element.rPr.rFonts.set(qn("w:eastAsia"), b.font)
            style.font.size = Pt(9.5)

    @staticmethod
    def _setup_page(section) -> None:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.55)
        section.bottom_margin = Cm(1.45)
        section.left_margin = Cm(1.65)
        section.right_margin = Cm(1.65)
        section.header_distance = Cm(0.55)
        section.footer_distance = Cm(0.55)

    def _asset(self, name: str) -> Path:
        for base in (self.assets, DEFAULT_ASSETS):
            if (base / name).is_file():
                return base / name
        raise FileNotFoundError(f"asset introuvable : {name} (cherché dans {self.assets} et {DEFAULT_ASSETS})")

    def _setup_header_footer(self, section) -> None:
        b = self.brand
        section.different_first_page_header_footer = True   # la couverture n'a ni en-tête ni pied

        header = section.header
        table = header.add_table(rows=1, cols=3, width=CONTENT_WIDTH)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        remove_table_borders(table)
        set_column_widths(table, [Cm(0.55), Cm(13.8), Cm(3.35)])
        icon = table.cell(0, 0).paragraphs[0]
        icon.add_run().add_picture(str(self._asset("camptocamp_icon.png")), width=Cm(0.30))
        set_text(table.cell(0, 1), b.kind, bold=True, color=b.accent, size=8.2, font=b.font)
        if b.client_logo:
            logo = table.cell(0, 2).paragraphs[0]
            logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            logo.add_run().add_picture(str(self._asset(b.client_logo)), width=b.client_logo_width)
        else:
            set_text(table.cell(0, 2), b.client, bold=True, color=DARK, size=8.2,
                     align=WD_ALIGN_PARAGRAPH.RIGHT, font=b.font)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        labels = LABELS.get(b.lang, LABELS["fr"])
        text = labels["copyright"].format(client=b.client, subject=b.subject or b.kind)
        style_run(p.add_run(text + "  "), size=8, bold=True, color=b.accent, font=b.font)
        add_field(p, "PAGE")
        style_run(p.add_run("/"), size=8, bold=True, color=DARK, font=b.font)
        add_field(p, "NUMPAGES")
        for run in p.runs:
            if run.font.size is None:
                style_run(run, size=8, bold=True, color=DARK, font=b.font)

    # -- briques ----------------------------------------------------------------
    def cover(self, title: str, subtitle: str, dateline: str, *, comments: str = "") -> None:
        """Couverture : logos, titre orange, sous-titre, ligne client + date. Pas de page de garde vide."""
        doc, b = self.doc, self.brand
        doc.core_properties.title = title
        doc.core_properties.subject = subtitle
        if comments:
            doc.core_properties.comments = comments
        logos = doc.add_table(rows=1, cols=2)
        logos.alignment = WD_TABLE_ALIGNMENT.CENTER
        logos.autofit = False
        remove_table_borders(logos)
        set_column_widths(logos, [Cm(12.4), Cm(4.8)])
        left = logos.cell(0, 0).paragraphs[0]
        left.add_run().add_picture(str(self._asset("camptocamp_logo.png")), width=Cm(9.5))
        if b.client_logo:
            right = logos.cell(0, 1).paragraphs[0]
            right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            right.add_run().add_picture(str(self._asset(b.client_logo)), width=b.client_logo_cover_width)
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(55)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        style_run(p.add_run(title), size=23, bold=True, color=b.accent, font=b.font)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        style_run(p.add_run(subtitle), size=17, bold=True, font=b.font)
        p = doc.add_paragraph()
        style_run(p.add_run(dateline), size=11.5, bold=True, font=b.font)
        doc.add_paragraph().paragraph_format.space_after = Pt(24)

    def heading(self, text: str, level: int = 1) -> None:
        p = self.doc.add_paragraph(style=f"Heading {level}")
        p.paragraph_format.keep_with_next = True
        p.add_run(text)

    def body(self, text: str, *, bold_lead: str | None = None, size: float = 9.5, italic: bool = False) -> None:
        """Paragraphe courant ; `bold_lead` met en gras le début de phrase qui porte l'idée."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.08
        f = self.brand.font
        if bold_lead and text.startswith(bold_lead):
            style_run(p.add_run(bold_lead), size=size, bold=True, font=f)
            style_run(p.add_run(text[len(bold_lead):]), size=size, italic=italic, font=f)
        else:
            style_run(p.add_run(text), size=size, italic=italic, font=f)

    def note(self, text: str) -> None:
        """Remarque en petit italique gris (provenance des captures, réserve, date de vérification)."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        style_run(p.add_run(text), size=8.5, italic=True, color=GREY, font=self.brand.font)

    def bullet(self, text: str, level: int = 0, color: str = DARK, size: float = 9.5) -> None:
        p = self.doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.55 + 0.45 * level)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        style_run(p.add_run(text), size=size, color=color, font=self.brand.font)

    def bullets(self, items) -> None:
        for item in items:
            self.bullet(item)

    def step(self, number: int, text: str, size: float = 9.5) -> None:
        """Étape numérotée : pastille orange + texte. Une action par étape, verbe à l'impératif."""
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        remove_table_borders(table)
        set_column_widths(table, [Cm(0.7), Cm(16.4)])
        f = self.brand.font
        set_text(table.cell(0, 0), str(number), bold=True, color=WHITE, size=9,
                 align=WD_ALIGN_PARAGRAPH.CENTER, font=f)
        set_cell_shading(table.cell(0, 0), self.brand.accent)
        set_text(table.cell(0, 1), text, size=size, font=f)
        for cell in table.rows[0].cells:
            set_cell_margins(cell, top=50, bottom=50)

    def steps(self, items) -> None:
        for number, text in enumerate(items, start=1):
            self.step(number, text)

    def callout(self, body: str, title: str = "", fill: str = LIGHT_ORANGE) -> None:
        """Encadré pâle : synthèse en tête de guide, garde-fou, décision, « vérifié le… »."""
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Cm(17.0)
        remove_table_borders(table)
        cell = table.cell(0, 0)
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        f = self.brand.font
        if title:
            style_run(p.add_run(title + " "), size=9.2, bold=True, color=self.brand.accent, font=f)
        style_run(p.add_run(body), size=9.2, color=DARK, font=f)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(0)

    def caption(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(5)
        style_run(p.add_run(text), size=8.5, italic=True, color=GREY, font=self.brand.font)

    def screenshot(self, filename: str, caption: str, width=SHOT_WIDTH) -> None:
        """Capture centrée + légende. Largeur pleine par défaut, `SHOT_WIDTH_NARROW` pour un popup."""
        path = self.shots / filename
        if not path.is_file():
            raise FileNotFoundError(f"capture absente : {path}")
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(str(path), width=width)
        self.caption(caption)

    def screenshot_pair(self, items, width=SHOT_WIDTH_PAIR) -> None:
        """Deux captures côte à côte, `items = [(fichier, légende), (fichier, légende)]` (avant/après)."""
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        remove_table_borders(table)
        set_column_widths(table, [Cm(8.55), Cm(8.55)])
        for cell, (filename, caption) in zip(table.rows[0].cells, items):
            set_cell_margins(cell, top=20, start=50, bottom=20, end=50)
            picture = cell.paragraphs[0]
            picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
            picture.paragraph_format.space_after = Pt(2)
            picture.add_run().add_picture(str(self.shots / filename), width=width)
            label = cell.add_paragraph()
            label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label.paragraph_format.space_after = Pt(3)
            style_run(label.add_run(caption), size=8.2, italic=True, color=GREY, font=self.brand.font)

    def matrix(self, headers, rows, widths=None, font_size: float = 8.2) -> None:
        """Tableau à en-tête orange et lignes zébrées : « ce que vous avez observé / ce que nous
        avons changé / statut », « décision / effet / effort », « bouton / ce qu'il fait »…
        La première colonne est en gras. Les lignes ne se coupent pas entre deux pages."""
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_borders(table)
        f = self.brand.font
        for idx, header in enumerate(headers):
            set_text(table.cell(0, idx), header, bold=True, color=WHITE, size=font_size, font=f)
            set_cell_shading(table.cell(0, idx), self.brand.accent)
        set_repeat_table_header(table.rows[0])
        for row_idx, values in enumerate(rows, start=1):
            cells = table.add_row().cells
            prevent_row_split(table.rows[-1])
            fill = WHITE if row_idx % 2 else LIGHT_GREY
            for idx, value in enumerate(values):
                set_text(cells[idx], value, bold=(idx == 0), size=font_size, font=f)
                set_cell_shading(cells[idx], fill)
        if widths:
            set_column_widths(table, widths)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(0)

    def spacer(self, points: float = 0) -> None:
        self.doc.add_paragraph().paragraph_format.space_after = Pt(points)

    def page_break(self) -> None:
        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # -- sortie ------------------------------------------------------------------
    def save(self, path: Path, *, pdf: bool = True) -> list[Path]:
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(path)
        out = [path]
        if pdf:
            out.append(to_pdf(path))
        return out


# --- Conversion et relecture -------------------------------------------------
def to_pdf(docx_path: Path) -> Path:
    """DOCX → PDF par LibreOffice headless. Le PDF est la version envoyée au client :
    le relire page par page (`pdf_pages_to_png`) avant de livrer."""
    docx_path = Path(docx_path)
    pdf_path = docx_path.with_suffix(".pdf")
    if pdf_path.exists():
        pdf_path.unlink()
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("LibreOffice introuvable (soffice) : impossible de produire le PDF")
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
        check=True, capture_output=True, text=True,
    )
    if not pdf_path.is_file():
        raise RuntimeError(f"la conversion n'a pas produit {pdf_path}")
    return pdf_path


def pdf_pages_to_png(pdf_path: Path, out_dir: Path, dpi: int = 70) -> list[Path]:
    """Une image par page, pour relire la mise en page (page blanche, capture coupée, tableau
    débordant) sans ouvrir de visionneuse."""
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    prefix = out_dir / Path(pdf_path).stem
    subprocess.run(["pdftoppm", "-r", str(dpi), "-png", str(pdf_path), str(prefix)], check=True)
    return sorted(out_dir.glob(f"{prefix.name}-*.png"))


def pdf_page_count(pdf_path: Path) -> int:
    out = subprocess.run(["pdfinfo", str(pdf_path)], check=True, capture_output=True, text=True).stdout
    for line in out.splitlines():
        if line.startswith("Pages:"):
            return int(line.split()[1])
    raise RuntimeError("pdfinfo : nombre de pages introuvable")
